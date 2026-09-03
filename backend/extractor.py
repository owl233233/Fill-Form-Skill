"""信息提取模块：从各类文件中提取结构化键值对信息"""
import os
import re

from openpyxl import load_workbook
from docx import Document

# ---------------- 字段别名归一化 ----------------
# 同一字段的多种写法归一化为标准键，保证资料库与模板占位符可以互相匹配
FIELD_ALIASES = {
    "姓名": ["name", "名字", "姓名name", "full_name", "fullname", "your name", "申请人", "联系人"],
    "电话": ["手机", "手机号", "手机号码", "联系电话", "联系方式", "电话号码", "phone", "tel", "mobile", "telephone", "phonenumber", "contact", "电话tel"],
    "邮箱": ["电子邮件", "电子邮箱", "email", "e-mail", "mail", "emailaddress", "邮箱email"],
    "公司": ["单位", "公司名称", "企业名称", "所在单位", "company", "organization", "org", "companyname", "employer"],
    "部门": ["院系", "所属部门", "department", "dept"],
    "职位": ["职务", "岗位", "职称", "title", "position", "job", "jobtitle"],
    "地址": ["联系地址", "通讯地址", "通信地址", "所在地", "住址", "家庭地址",
             "家庭住址", "居住地址", "现住址", "现居住地", "户籍地址",
             "address", "地址address"],
    "性别": ["gender", "sex"],
    "年龄": ["age"],
    "学历": ["教育程度", "文化程度", "education", "degree"],
    "专业": ["major", "专业名称"],
    "学校": ["毕业院校", "院校", "university", "school", "college"],
    "日期": ["date", "填表日期", "申请日期", "提交日期", "时间"],
    "身份证号": ["身份证", "证件号码", "id", "idcard", "idnumber", "身份证号码"],
    "学号": ["工号", "编号", "员工编号"],
    "备注": ["note", "notes", "remark", "说明"],
}

_ALIAS_INDEX = {}
for std, aliases in FIELD_ALIASES.items():
    _ALIAS_INDEX[std] = std
    for a in aliases:
        _ALIAS_INDEX[a] = std


def normalize_key(key):
    """键名归一化：去空格、转小写，再查别名表；未命中别名表则返回清理后的原始键"""
    k = re.sub(r"[\s\u3000_\-·、，,：:（）()【】\[\]]+", "", str(key)).lower()
    if not k:
        return None
    return _ALIAS_INDEX.get(k, k)


def image_name_key(filename):
    """图片素材名归一化：去扩展名、去重名序号后缀（头像_1.png → 头像），再归一化"""
    stem = os.path.splitext(os.path.basename(filename))[0]
    stem = re.sub(r"[_\-]\d{1,3}$", "", stem)
    return normalize_key(stem) or stem.lower()


def lookup_field(field_name, entries):
    """
    智能查找字段：entries 为 {标准键: 条目}。
    ① 精确匹配（归一化后相等 / 命中别名表）
    ② 后缀匹配：资料键以字段名结尾或字段名以资料键结尾（如「现居地址」↔「地址」、「入职部门」↔「部门」）
    返回 (命中的资料键, 条目) 或 None。
    """
    nk = normalize_key(field_name)
    if not nk:
        return None
    if nk in entries:
        return nk, entries[nk]
    candidates = []
    for k in entries:
        if len(k) < 2 or len(nk) < 2:
            continue
        if k.endswith(nk) and len(k) > len(nk):
            candidates.append((len(k) - len(nk), k))
        elif nk.endswith(k) and len(nk) > len(k):
            candidates.append((len(nk) - len(k), k))
    if candidates:
        candidates.sort(key=lambda x: (x[0], len(x[1])))  # 差距最小 = 最精确
        k = candidates[0][1]
        return k, entries[k]
    return None


# ---------------- 通用文本键值提取 ----------------
# 匹配 "标签：值" / "标签: 值" / "标签：值（中文全角冒号）" / "标签=value"
# 标签里允许含空格（如简历里的「姓    名：张三」「住    址：…」）
_KV_PATTERN = re.compile(
    r"^[丨|●•\s]*([\u4e00-\u9fa5A-Za-z0-9（）()\s]{1,20})\s*[:：=]\s*(.+?)\s*$"
)
# 行内提取（一行出现多个键值时按分隔符切分）
_INLINE_SPLIT = re.compile(r"[;；\t]{1,}")

PHONE_RE = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
IDCARD_RE = re.compile(r"(?<!\d)\d{17}[0-9Xx](?!\d)")
URL_RE = re.compile(r"https?://[^\s，。；,;]+")

# 键名里包含这些词的倾向于"标签"而不是正文
_LABEL_HINT = re.compile(r"姓名|名字|电话|手机|联系|邮箱|邮件|公司|单位|部门|职位|职务|岗位|职称|地址|性别|年龄|学历|专业|院校|学校|日期|时间|身份证|证件|学号|工号|编号|邮编|传真|微信|QQ|备注|说明|标题|名称|简介|描述")


def extract_kv_from_text(text):
    """从纯文本提取键值对。返回 {标准键: {"key": 原始键, "value": 值}}"""
    result = {}
    if not text:
        return result

    def put(raw_key, value):
        value = str(value).strip().strip("；;，,。.、").strip()
        if not value or len(value) > 200:
            return
        nk = normalize_key(raw_key)
        if not nk:
            return
        # 已存在时保留更短的值（往往是更干净的提取结果），相同则跳过
        if nk in result:
            if len(value) < len(result[nk]["value"]):
                result[nk] = {"key": str(raw_key).strip(), "value": value}
        else:
            result[nk] = {"key": str(raw_key).strip(), "value": value}

    for line in text.splitlines():
        line = line.strip()
        if not line or len(line) > 500:
            continue
        # 行内可能含多个键值（用分号/制表符分隔）
        for seg in _INLINE_SPLIT.split(line):
            seg = seg.strip()
            if not seg:
                continue
            m = _KV_PATTERN.match(seg)
            if m:
                raw_key, value = m.group(1).strip(), m.group(2).strip()
                # 值里如果还有冒号结构（如"邮箱：a@b.com：备注"）取第一段
                if _LABEL_HINT.search(raw_key) or len(raw_key) <= 8:
                    put(raw_key, value)

    # 模式识别兜底：即便没有"标签："结构，也识别电话/邮箱/身份证
    phones = PHONE_RE.findall(text)
    if phones and "电话" not in result:
        put("电话", phones[0])
    emails = EMAIL_RE.findall(text)
    if emails and "邮箱" not in result:
        put("邮箱", emails[0])
    idcards = IDCARD_RE.findall(text)
    if idcards and "身份证号" not in result:
        put("身份证号", idcards[0])
    return result


# ---------------- 各格式文本抽取 ----------------

def _iter_docx_paragraphs(doc):
    for p in doc.paragraphs:
        yield p.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text
    # 文本框（textbox）内容：python-docx 的 paragraphs / tables 都不含文本框，
    # 而很多简历/海报/申报表的正文就放在文本框里。从 XML 层补抓。
    try:
        from docx.oxml.ns import qn
        for txbx in doc.element.body.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                line = "".join(t.text or "" for t in p.iter(qn("w:t")))
                if line.strip():
                    yield line
    except Exception:
        pass


_MACRO_MAIN_CT = "application/vnd.ms-word.document.macroEnabled.main+xml"
_PLAIN_MAIN_CT = ("application/vnd.openxmlformats-officedocument"
                  ".wordprocessingml.document.main+xml")


def open_docx_any(path):
    """打开 .docx / .docm。

    python-docx 会按包内声明的 content type 校验，.docm 声明的是
    macroEnabled，因此直接 Document(path) 会抛
    "file ... is not a Word file"。仅改扩展名无效——声明在 zip 内部。

    这里在内存中把 [Content_Types].xml 里的声明改成标准值再加载，
    不改动用户的原文件。宏内容不会被读取（也不执行）。
    """
    try:
        return Document(path)
    except ValueError:
        pass
    import io
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(path) as zin:
        entries = [(i, zin.read(i.filename)) for i in zin.infolist()]
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in entries:
            if info.filename == "[Content_Types].xml":
                text = data.decode("utf-8", "ignore")
                data = text.replace(_MACRO_MAIN_CT, _PLAIN_MAIN_CT).encode("utf-8")
            zout.writestr(info, data)
    buf.seek(0)
    return Document(buf)


def extract_docx_text(path):
    try:
        doc = open_docx_any(path)
        return "\n".join(t for t in _iter_docx_paragraphs(doc) if t and t.strip())
    except Exception:
        return ""


def extract_xlsx_text(path):
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c).strip() for c in row]
                if any(cells):
                    lines.append(" | ".join(c for c in cells if c))
        wb.close()
        return "\n".join(lines)
    except Exception:
        return ""


def extract_pdf_text(path):
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:20]:
                t = page.extract_text() or ""
                if t:
                    texts.append(t)
        return "\n".join(texts)
    except Exception:
        return ""


def extract_txt_text(path):
    for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, ValueError):
            continue
    return ""


def extract_json_text(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return ""


# ---------------- Excel 键值专项提取 ----------------

def extract_xlsx_kv(path):
    """Excel 更可能是"键值表"或"表头+数据行"，做专项解析"""
    result = {}
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return result
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            # 情况1：两列结构（第一列像标签，值非空）
            two_col = all(len([c for c in r if c is not None and str(c).strip()]) <= 2 for r in rows if r)
            if two_col:
                for r in rows:
                    r = list(r) + [None] * 4
                    k, v = r[0], r[1]
                    if k is None or v is None:
                        continue
                    k, v = str(k).strip(), str(v).strip()
                    if k and v and len(k) <= 20:
                        nk = normalize_key(k)
                        if nk and nk not in result:
                            result[nk] = {"key": k, "value": v}
            else:
                # 情况2：表头+数据行，取第一行做表头
                header = [str(c).strip() if c is not None else "" for c in rows[0]]
                data_row = None
                for r in rows[1:]:
                    if any(c is not None and str(c).strip() for c in r):
                        data_row = r
                        break
                if data_row:
                    for h, v in zip(header, list(data_row) + [None] * len(header)):
                        if h and v is not None and len(h) <= 20:
                            nk = normalize_key(h)
                            if nk and nk not in result:
                                result[nk] = {"key": h, "value": str(v).strip()}
    finally:
        wb.close()
    return result


# ---------------- 统一入口 ----------------

TEXT_EXTRACTORS = {
    ".docx": extract_docx_text,
    ".docm": extract_docx_text,   # 带宏的 Word，用 open_docx_any 兼容
    ".xlsx": extract_xlsx_text,
    ".xls": extract_xlsx_text,
    ".txt": extract_txt_text,
    ".md": extract_txt_text,
    ".pdf": extract_pdf_text,
    ".json": extract_json_text,
    ".csv": extract_txt_text,
}

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}


def file_category(filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext in TEXT_EXTRACTORS:
        return "document"
    return "other"


def extract_info(path, filename):
    """提取入口：返回 (kv_dict, text_preview)"""
    ext = os.path.splitext(filename)[1].lower()
    kv = {}
    text = ""
    if ext == ".xlsx" or ext == ".xls":
        kv = extract_xlsx_kv(path)
        text = extract_xlsx_text(path)
    elif ext in TEXT_EXTRACTORS:
        text = TEXT_EXTRACTORS[ext](path) or ""
        kv = extract_kv_from_text(text)
    # 预览文本（最多 400 字）
    preview = re.sub(r"\n{2,}", "\n", text).strip()[:400]
    return kv, preview
