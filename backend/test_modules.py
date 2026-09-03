"""模块级测试：提取器 + 填充器"""
import os
import shutil
import sys

# 原为云端硬编码路径 /workspace/form-filler/... 与 /tmp/ff_test，本地不可用，改为相对定位
BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BACKEND_DIR)
sys.path.insert(0, BACKEND_DIR)
import extractor
import filler
import form_analyzer

SAMPLE = os.path.join(PROJECT_DIR, "sample")
TMP = os.path.join(PROJECT_DIR, "tmp", "ff_test")
os.makedirs(TMP, exist_ok=True)
ok, fail = 0, 0

def check(name, cond, detail=""):
    global ok, fail
    if cond:
        ok += 1
        print(f"  ✓ {name}")
    else:
        fail += 1
        print(f"  ✗ {name} {detail}")

print("== 1. 信息提取 ==")
kv, preview = extractor.extract_info(os.path.join(SAMPLE, "李云飞个人简历.docx"), "李云飞个人简历.docx")
print("  提取结果:", {v["key"]: v["value"] for v in kv.values()})
check("提取姓名", kv.get("姓名", {}).get("value") == "李云飞")
check("提取电话", kv.get("电话", {}).get("value") == "13812345678")
check("提取邮箱", kv.get("邮箱", {}).get("value") == "liyunfei@example.com")
check("提取学校", kv.get("学校", {}).get("value") == "复旦大学")
check("提取专业", kv.get("专业", {}).get("value") == "计算机科学与技术")

kv2, _ = extractor.extract_info(os.path.join(SAMPLE, "岗位信息.txt"), "岗位信息.txt")
r = extractor.lookup_field("部门", {**kv, **kv2}); check("智能匹配部门(入职部门→部门)", r is not None and r[1]["value"] == "人工智能事业部")
check("txt提取职位", kv2.get("职位", {}).get("value") == "算法工程师")
check("txt提取公司", kv2.get("公司", {}).get("value") == "智创未来科技有限公司")

print("== 2. 模板占位符解析 ==")
tf, imgs = filler.parse_docx_placeholders(os.path.join(SAMPLE, "员工入职登记表模板.docx"))
print("  文本字段:", sorted(tf))
print("  图片字段:", sorted(imgs))
check("docx解析含姓名", "姓名" in tf)
check("docx解析含电话", "电话" in tf)
check("docx解析含当前日期", "当前日期" in tf)
check("docx解析含页眉公司", "公司" in tf)
check("docx解析图片占位符", "头像" in imgs)

tf2, imgs2 = filler.parse_xlsx_placeholders(os.path.join(SAMPLE, "产品报价单模板.xlsx"))
print("  xlsx文本字段:", sorted(tf2))
check("xlsx解析含姓名", "姓名" in tf2)
check("xlsx解析含产品名称", "产品名称" in tf2)

print("== 3. docx 填充（跨 run + 格式保持 + 图片） ==")
# 模拟 API 层组装：对每个模板字段做智能匹配，统一以模板字段归一化键存放
values = {}
for name in tf:
    hit = extractor.lookup_field(name, {**kv, **kv2})
    if hit:
        values[filler.normalize(name)] = hit[1]["value"]
values.update({"性别": "男", "签名": "李云飞", "签署日期": "2026-08-29"})
img_map = {"头像": os.path.join(SAMPLE, "头像.png")}
out_docx = os.path.join(TMP, "output_员工入职登记表.docx")
res = filler.fill_docx(os.path.join(SAMPLE, "员工入职登记表模板.docx"), values, img_map, out_docx)
print("  填充结果:", res)
check("图片插入成功", res["images_inserted"] == 1)
check("无缺失图片", res["images_missing"] == [])

# 验证生成文档内容与格式
from docx import Document
d = Document(out_docx)
all_text = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += "\n" + cell.text
header_text = d.sections[0].header.paragraphs[0].text
check("姓名已填入", "李云飞" in all_text)
check("电话已填入", "13812345678" in all_text)
check("动态日期已填入", "2026-" in all_text)
check("无残留占位符", "{{" not in all_text and "[[" not in all_text)
check("页眉公司已填入", "智创未来科技有限公司" in header_text)
check("文档含图片", len(d.inline_shapes) == 1)

# 格式验证：标题仍为 22pt 加粗深蓝
h = d.paragraphs[0]
r = h.runs[0]
check("标题字号保持22pt", r.font.size.pt == 22, f"实际{r.font.size.pt if r.font.size else 'None'}")
check("标题加粗保持", r.font.bold is True)
check("标题颜色保持", str(r.font.color.rgb) == "1F3A8F")

# 验证图片占位符段落
img_para_texts = [p.text for p in d.paragraphs if "员工照片" in p.text]
check("图片占位符文本已清除", all("{{" not in t for t in img_para_texts))

print("== 4. xlsx 填充 ==")
xvalues = dict(values)
xvalues.update({"产品名称": "智能填表系统企业版", "数量": "10", "单价": "8800", "备注": "含一年技术支持"})
out_xlsx = os.path.join(TMP, "output_报价单.xlsx")
res2 = filler.fill_xlsx(os.path.join(SAMPLE, "产品报价单模板.xlsx"), xvalues, {}, out_xlsx)
from openpyxl import load_workbook
wb = load_workbook(out_xlsx)
ws = wb.active
check("xlsx公司已填入", "智创未来科技有限公司" in str(ws["A1"].value))
check("xlsx产品已填入", "智能填表系统企业版" in str(ws["A6"].value))
check("xlsx标题样式保持", ws["A1"].font.size == 16 and ws["A1"].font.bold)
check("xlsx表头填充色保持", ws["A5"].fill.fgColor.rgb in ("004F46E5", "FF4F46E5"))
check("xlsx无残留占位符", all("{{" not in str(c.value) for row in ws.iter_rows() for c in row if c.value))

print("== 5. 命名规则渲染 ==")
name1 = filler.render_filename("{{姓名}}_入职登记表_{{当前日期}}", "模板.docx", values, ".docx")
print("  渲染:", name1)
check("命名含姓名", name1.startswith("李云飞_入职登记表_2026"))
check("命名含扩展名", name1.endswith(".docx"))
name2 = filler.render_filename("", "员工入职登记表模板.docx", values, ".docx")
print("  默认命名:", name2)
check("默认命名规则", name2.startswith("员工入职登记表模板_"))
name3 = filler.render_filename("{{姓名}}_报告", "t.docx", {}, ".docx")
print("  缺值命名:", name3)
check("缺失字段留空但不报错", name3 == "_报告.docx" or "报告" in name3)

print("== 6. 论文引用解析 ==")
cit = form_analyzer.parse_citation(
    'Lee, H., Zhao, K. and Zou, F. (2022) "Does the Early Retirement Policy Really '
    'Benefit Women?" Journal of Economic Behavior and Organization. (196) 330-345. '
    'SSCI Q2，中科院3区，影响因子2.6. 文章引用10次', self_hints=["Zou"])
check("解析论文1标题", cit["title"] == "Does the Early Retirement Policy Really Benefit Women?")
check("解析论文1期刊", cit["journal"] == "Journal of Economic Behavior and Organization")
check("解析论文1ISSN", cit["issn"] == "0167-2681")
check("解析论文1影响因子", cit["impact_factor"] == "2.6")
check("解析论文1收录", "SSCI" in cit["indexing"] and "中科院3区" in cit["indexing"])
check("解析论文1年卷页", cit["year_vol_pages"] == "2022, 196: 330-345")
check("解析论文1作者排序(第三作者)", cit["rank"] == "第三作者")
check("解析论文1引用次数", cit["cited"] == "10")

cit2 = form_analyzer.parse_citation(
    'Zou, F. and Shi, L. (2026) "Changes in the Life-Cycle Earnings of Chinese Workers: '
    'The Role of Education, Gender and Economic Reform" Post Communist Economies，1-28，'
    'SSCI Q2, 中科院3区。影响因子2.0', self_hints=["Zou"])
check("解析论文2作者排序(第一作者)", cit2["rank"] == "第一作者")
check("解析论文2ISSN", cit2["issn"] == "1463-1377")
check("解析论文2影响因子", cit2["impact_factor"] == "2.0")
check("解析论文2年卷页(无卷)", cit2["year_vol_pages"] == "2026: 1-28")

sec = form_analyzer.extract_papers_from_text(
    '发表论文\n'
    'Lee, H., Zhao, K. and Zou, F. (2022) "T1" Journal of Economic Behavior and '
    'Organization. (196) 330-345. SSCI Q2，中科院3区，影响因子2.6\n'
    'Zou, F. and Shi, L. (2026) "T2" Post Communist Economies，1-28，SSCI Q2, 中科院3区。影响因子2.0\n'
    '工作论文\n', self_hints=["Zou"])
check("文本抽论文2条", len(sec.get("发表论文", {}).get("rows", [])) == 2)
check("无星号判通讯作者留空", cit["corresponding"] == "" and cit2["corresponding"] == "")

print(f"\n结果：通过 {ok} 项，失败 {fail} 项")
sys.exit(1 if fail else 0)
