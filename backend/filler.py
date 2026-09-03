"""模板解析与填充模块：解析 {{字段}} 占位符，填充时严格保持模板原有格式"""
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Emu
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from PIL import Image as PILImage

from extractor import normalize_key  # 与资料库统一使用同一套字段归一化（含别名表）


def normalize(s):
    """字段归一化：与 extractor 保持一致（别名映射），保证模板占位符能命中资料库键"""
    return normalize_key(s) or str(s).strip().lower()

# 占位符：{{字段}} 为主；同时兼容 [[字段]]（部分办公软件花括号输入不便）
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}<>]+?)\s*\}\}|\[\[\s*([^\[\]<>]+?)\s*\]\]")
# 图片占位符：{{img:名称}} / {{图片:名称}}
IMG_PLACEHOLDER_RE = re.compile(r"^\s*(?:img|image|图片|照片)\s*[:：]\s*(.+)$", re.IGNORECASE)
# 动态字段（无需资料库提供）
DYNAMIC_FIELDS = {"当前日期", "当前时间", "today", "now"}


def _placeholder_name(match_text):
    """去掉 {{}} 拿到字段名；若是图片占位符返回 ('img', 名称)，否则 ('text', 字段名)"""
    name = match_text.strip()
    m = IMG_PLACEHOLDER_RE.match(name)
    if m:
        return "img", m.group(1).strip()
    return "text", name


def dynamic_value(field):
    """返回动态字段的当前值（当前日期 / 当前时间）。不是动态字段返回 None。"""
    now = datetime.now()
    if normalize(field) in {"当前日期", "today"}:
        return now.strftime("%Y-%m-%d")
    if normalize(field) in {"当前时间", "now"}:
        return now.strftime("%Y-%m-%d %H:%M")
    return None


# ============================================================
#                     DOCX 解析与填充
# ============================================================

def _iter_docx_paragraphs_all(doc):
    """正文段落 + 表格内段落 + 页眉页脚段落"""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer, section.first_page_header,
                     section.first_page_footer, section.even_page_header, section.even_page_footer):
            try:
                for p in part.paragraphs:
                    yield p
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
            except Exception:
                continue


def parse_docx_placeholders(path):
    """解析 docx 模板中的所有占位符，返回 (text_fields:set, img_fields:set)"""
    doc = Document(path)
    text_fields, img_fields = set(), set()
    for p in _iter_docx_paragraphs_all(doc):
        for m in PLACEHOLDER_RE.finditer(p.text):
            kind, name = _placeholder_name(m.group(1) or m.group(2))
            (img_fields if kind == "img" else text_fields).add(name)
    return text_fields, img_fields


def _replace_runs(paragraph, text_values, img_sources, used_images):
    """在单个段落内做跨 run 占位符替换，保持占位符起始 run 的格式"""
    runs = paragraph.runs
    if not runs:
        return
    texts = [r.text or "" for r in runs]
    combined = "".join(texts)
    matches = list(PLACEHOLDER_RE.finditer(combined))
    if not matches:
        return

    # 从后往前处理，避免前面修改导致后面偏移失效
    for m in reversed(matches):
        kind, name = _placeholder_name(m.group(1) or m.group(2))
        start, end = m.span()
        dyn = dynamic_value(name)
        if kind == "text":
            value = text_values.get(normalize(name))
            if value is None and dyn is not None:
                value = dyn
            if value is None:
                continue  # 没有对应值，保留原占位符
            value = str(value)
        else:
            value = ""  # 图片占位符文本替换为空

        # 定位占位符跨越的 runs
        pos = 0
        affected = []
        for i, t in enumerate(texts):
            rs, re_ = pos, pos + len(t)
            pos = re_
            if re_ <= start or rs >= end:
                continue
            affected.append((i, rs, re_))
        if not affected:
            continue

        i0, rs0, _ = affected[0]
        il, rsl, _ = affected[-1]
        before = texts[i0][: start - rs0]
        after = texts[il][end - rsl:]

        # 中间 run 文本清空
        for i, _, _ in affected:
            texts[i] = ""

        if kind == "text":
            if il == i0:
                texts[i0] = before + value + after
            else:
                texts[i0] = before + value
                texts[il] = after
        else:
            # 图片：文本部分合并，图片插入到第一个受影响 run 处
            if il == i0:
                texts[i0] = before + after
            else:
                texts[i0] = before
                texts[il] = after
            used_images.append((runs[i0], name))

    # 写回 run 文本
    for run, t in zip(runs, texts):
        if run.text != t:
            run.text = t


def _insert_picture(run, image_path, doc):
    """在指定 run 处插入内联图片（等比缩放，不超过版心宽度），保持所在位置样式"""
    try:
        section = doc.sections[0]
        max_w = section.page_width - section.left_margin - section.right_margin
    except Exception:
        max_w = Emu(6096000)  # A4 约 6.35 英寸可用宽度兜底
    try:
        with PILImage.open(image_path) as im:
            w_px, h_px = im.size
            dpi = im.info.get("dpi", (96, 96))[0] or 96
            w_emu = int(w_px / dpi * 914400)
            h_emu = int(h_px / dpi * 914400)
        if w_emu <= 0 or h_emu <= 0:
            raise ValueError
    except Exception:
        w_emu = h_emu = None

    try:
        if w_emu and w_emu > max_w:
            ratio = max_w / w_emu
            run.add_picture(image_path, width=int(w_emu * ratio), height=int(h_emu * ratio))
        elif w_emu:
            run.add_picture(image_path, width=w_emu, height=h_emu)
        else:
            run.add_picture(image_path)
        return True
    except Exception:
        return False


def fill_docx(template_path, text_values, image_map, output_path):
    """
    填充 Word 模板。
    text_values: {标准化字段名: 值}
    image_map:   {标准化图片名: 图片路径}
    """
    doc = Document(template_path)
    text_values = {normalize(k): v for k, v in (text_values or {}).items()}
    image_map = {normalize(k): v for k, v in (image_map or {}).items()}

    inserted, missed = 0, []
    for p in _iter_docx_paragraphs_all(doc):
        used_images = []
        _replace_runs(p, text_values, image_map, used_images)
        for run, img_name in used_images:
            src = image_map.get(normalize(img_name))
            if src and os.path.exists(src):
                if _insert_picture(run, src, doc):
                    inserted += 1
            else:
                missed.append(img_name)
    doc.save(output_path)
    return {"images_inserted": inserted, "images_missing": sorted(set(missed))}


# ============================================================
#                     XLSX 解析与填充
# ============================================================

def parse_xlsx_placeholders(path):
    wb = load_workbook(path, data_only=False)
    text_fields, img_fields = set(), set()
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    for m in PLACEHOLDER_RE.finditer(cell.value):
                        kind, name = _placeholder_name(m.group(1) or m.group(2))
                        (img_fields if kind == "img" else text_fields).add(name)
    wb.close()
    return text_fields, img_fields


def fill_xlsx(template_path, text_values, image_map, output_path):
    wb = load_workbook(template_path, data_only=False)
    text_values = {normalize(k): v for k, v in (text_values or {}).items()}
    image_map = {normalize(k): v for k, v in (image_map or {}).items()}

    inserted, missed = 0, []
    for ws in wb.worksheets:
        img_anchors = []
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                val = cell.value
                if not PLACEHOLDER_RE.search(val):
                    continue

                def _sub(m):
                    kind, name = _placeholder_name(m.group(1) or m.group(2))
                    if kind == "img":
                        img_anchors.append((cell.coordinate, name))
                        return ""
                    v = text_values.get(normalize(name))
                    if v is None:
                        dyn = dynamic_value(name)
                        if dyn is not None:
                            v = dyn
                    return str(v) if v is not None else m.group(0)

                new_val = PLACEHOLDER_RE.sub(_sub, val)
                if new_val != val:
                    cell.value = new_val if new_val != "" else None

        for coord, img_name in img_anchors:
            src = image_map.get(normalize(img_name))
            if src and os.path.exists(src):
                try:
                    img = XLImage(src)
                    img.anchor = coord
                    ws.add_image(img)
                    inserted += 1
                except Exception:
                    pass
            else:
                missed.append(img_name)
    wb.save(output_path)
    return {"images_inserted": inserted, "images_missing": sorted(set(missed))}


# ============================================================
#                     统一入口 & 命名渲染
# ============================================================

SUPPORTED_TEMPLATE_EXTS = {".docx", ".xlsx"}


def parse_template(path, filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return parse_docx_placeholders(path)
    if ext == ".xlsx":
        return parse_xlsx_placeholders(path)
    raise ValueError(f"不支持的模板格式：{ext}（目前支持 .docx / .xlsx）")


def fill_template(path, filename, text_values, image_map, output_path):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return fill_docx(path, text_values, image_map, output_path)
    if ext == ".xlsx":
        return fill_xlsx(path, text_values, image_map, output_path)
    raise ValueError(f"不支持的模板格式：{ext}")


_INVALID_FS_CHARS = re.compile(r'[\\/:*?"<>|\r\n\t]')


def render_filename(pattern, template_name, text_values, ext):
    """
    按模板命名规则渲染输出文件名。
    pattern 中的 {{字段}} 用资料值替换；无 pattern 时用 "模板名_当前日期"。
    """
    values = {normalize(k): str(v) for k, v in (text_values or {}).items()}

    def _sub(m):
        kind, name = _placeholder_name(m.group(1) or m.group(2))
        if kind == "img":
            return ""
        v = values.get(normalize(name))
        if not v:
            dyn = dynamic_value(name)
            if dyn is not None:
                v = dyn
        return str(v).strip() if v else ""

    if pattern and pattern.strip():
        name = PLACEHOLDER_RE.sub(_sub, pattern.strip())
        # 命名规则里可能自带扩展名
        if not os.path.splitext(name)[1]:
            name = f"{name}{ext}"
    else:
        base = os.path.splitext(template_name)[0]
        name = f"{base}_{datetime.now().strftime('%Y%m%d')}{ext}"

    name = _INVALID_FS_CHARS.sub("_", name).strip("_ ")
    if not name or name == ext:
        name = f"filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
    # 防止文件名超长
    stem, e = os.path.splitext(name)
    if len(stem) > 80:
        name = stem[:80] + e
    return name
