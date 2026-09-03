# -*- coding: utf-8 -*-
"""表格型表单适配：识别并填充「标签格 + 空格」形式的传统表单。

背景
----
很多单位/高校的表格（申报书、登记表、信息表）**不是** `{{占位符}}` 模板，
而是「标签 + 待填空格」，例如：

    表0  r0  c0=[姓    名：]  c1=(空)
    表1  r0  c0=[姓  名] c1=(空) | c3=[性别] c4=(空)

这类表格里一个 `{}` 都没有，原有的 `filler.parse_template` 只能识别到 0 个字段。

设计：模板驱动
--------------
不是盲目解析资料再瞎匹配，而是：

1. 先分析**模板**，得到它要哪些字段（标签名 + 待填格坐标）
2. 再拿这些标签去**资料**里找同名（或别名/后缀匹配）的标签格，取其相邻格的值
3. 最后按坐标写回

这样能避开「表头行被当成数据」「合并单元格重复读取」等一堆坑，
而且只查需要的字段，资料里的无关内容一概不碰。

字段类型
--------
- `cell_label`：标签格 + 右侧空值格
- `example`   ：标签格 + 右侧示例文本（`XXX学校XXX学博士`）
- `inline`    ：段落内 `标签：____`
- `placeholder`：`{{name}}`（沿用原有机制）
"""
import copy
import re

from docx.text.paragraph import Paragraph

# 示例占位文本：XXX学校、X年X月、XX学学士
_EXAMPLE_PAT = re.compile(r"[Xx]{2,}")
# 纯日期/编号之类，不可能当标签
_NOT_LABEL_PAT = re.compile(r"^[\d\s./年月日—\-~至]+$")

# 段落内的行内字段：`课程名称：______`
_INLINE_RE = re.compile(
    r"^\s*([\u4e00-\u9fa5A-Za-z0-9（）()]{1,14})\s*[：:]\s*([_＿\s.。·…]*)$")


# ---------------------------------------------------------------- 基础工具

def norm_label(text):
    """'姓    名：' -> '姓名'；去空白、去尾部冒号、去前导序号"""
    t = re.sub(r"[\s\u3000]+", "", text or "")
    t = t.rstrip("：:")
    t = re.sub(r"^[（(]?\d{1,2}[)）、.]\s*", "", t)
    return t


def is_example_text(text):
    """是否形如 `XXX学校` / `X年X月` 的示例占位文本。

    注意不能用 `[Xx]{2,}`（要求 X 连续）——`X年X月至X年X月`、`X/X`
    这类 X 是分开的，照样是示例文本。按 X 的总个数判断更可靠。
    """
    t = (text or "").strip()
    if not t:
        return False
    return len(_EXAMPLE_CHARS.findall(t)) >= 2


_EXAMPLE_CHARS = re.compile(r"[Xx]")


def iter_row_cells(row):
    """遍历一行中"真实"的单元格，跳过合并单元格的重复引用。

    python-docx 的 `row.cells` 对合并单元格会重复返回同一底层元素，
    不做去重的话统计和取值都会错乱（实测"财经学院"被读了 13 次）。

    返回 [(grid_col_index, cell), ...]
    """
    seen, out = set(), []
    for ci, c in enumerate(row.cells):
        if id(c._tc) in seen:
            continue
        seen.add(id(c._tc))
        out.append((ci, c))
    return out


def _is_label_like(text):
    """像不像一个标签（而不是一段正文或一个值）"""
    t = (text or "").strip()
    if not t or len(t) > 26:
        return False
    if _NOT_LABEL_PAT.match(t):
        return False
    if any(p in t for p in ("。", "；", "，", "、")):
        return False
    if is_example_text(t):
        return False
    return True


def _row_texts(row):
    return [c.text.strip() for _, c in iter_row_cells(row)]


def _is_section_title(row, min_span_ratio=0.55):
    """是否是分节标题：整行只有一个非空单元格，且它横跨该行大部分列宽。

    只看「非空单元格数 == 1」会误判——学习经历表里 `[大专]` 这种
    只有阶段名、其余全空的数据行也会被当成标题（实测就出过这个错，
    结果学习经历被拦腰截断、丢掉了本科那一行）。

    分节标题的单元格通常是横向合并、横跨整行的，据此区分。
    """
    filled = [(ci, c) for ci, c in iter_row_cells(row) if c.text.strip()]
    if len(filled) != 1:
        return False
    cell = filled[0][1]
    raw = row.cells
    span = sum(1 for x in raw if x._tc is cell._tc)
    return span >= max(2, len(raw) * min_span_ratio)


# ---------------------------------------------------------------- 模板分析

def analyze_template_fields(doc):
    """识别模板中的待填字段。

    返回 [ {kind, name, table, row, col | para, current}, ... ]
    """
    fields = []

    for ti, tb in enumerate(doc.tables):
        for ri, row in enumerate(tb.rows):
            cells = iter_row_cells(row)
            for i, (ci, cell) in enumerate(cells):
                text = cell.text.strip()
                if not text or not _is_label_like(text):
                    continue
                label = norm_label(text)
                if not label:
                    continue
                if i + 1 >= len(cells):
                    continue
                nci, ncell = cells[i + 1]
                ntext = ncell.text.strip()
                if ntext == "":
                    fields.append({"kind": "cell_label", "name": label,
                                   "table": ti, "row": ri, "col": nci, "current": ""})
                elif is_example_text(ntext):
                    fields.append({"kind": "example", "name": label,
                                   "table": ti, "row": ri, "col": nci, "current": ntext})
                # 值格已有真实内容（既非空也非示例）→ 视为已填，不动

    for pi, p in enumerate(doc.paragraphs):
        m = _INLINE_RE.match(p.text or "")
        if m:
            fields.append({"kind": "inline", "name": norm_label(m.group(1)),
                           "para": pi, "current": m.group(2)})

    return fields


# ---------------------------------------------------------------- 资料索引

def build_source_index(doc):
    """从资料文档建「归一化标签 -> 值」的索引。

    只按「标签格 + 值格」交替配对取值，并跳过列表区块的表头行
    （表头行上方通常是独占一行的分节标题，如 [学习经历]）。
    """
    index = {}
    for tb in doc.tables:
        rows = tb.rows
        for ri, row in enumerate(rows):
            # 分节标题的下一行是表头，表头不是数据，跳过
            if ri > 0 and _is_section_title(rows[ri - 1]):
                continue
            cells = iter_row_cells(row)
            if len(cells) < 2:
                continue
            for i in range(0, len(cells) - 1, 2):
                lt = norm_label(cells[i][1].text)
                vt = cells[i + 1][1].text.strip()
                if not lt or not vt:
                    continue
                if is_example_text(lt) or is_example_text(vt):
                    continue
                index.setdefault(lt, vt)
    return index


# ---------------------------------------------------------------- 填充

def set_cell_text(cell, text):
    """写入单元格，保留首个 run 的字体格式；多行按行追加段落。"""
    lines = str(text).split("\n")
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run("")
    p.runs[0].text = lines[0]
    for r in p.runs[1:]:
        r.text = ""
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    prev = p._element
    for line in lines[1:]:
        newp = copy.deepcopy(p._element)
        prev.addnext(newp)
        prev = newp
        para = Paragraph(newp, p._parent)
        para.runs[0].text = line
        for r in para.runs[1:]:
            r.text = ""


def set_paragraph_text(p, text):
    """改写段落文本，保留首个 run 的格式。"""
    if not p.runs:
        p.add_run("")
    p.runs[0].text = str(text)
    for r in p.runs[1:]:
        r.text = ""


def fill_fields(doc, fields, values):
    """按字段坐标把 values 写回模板。返回实际写入的字段列表。"""
    written = []
    for f in fields:
        val = values.get(f["name"])
        if not val:
            continue
        try:
            if f["kind"] in ("cell_label", "example"):
                set_cell_text(doc.tables[f["table"]].rows[f["row"]].cells[f["col"]], val)
            elif f["kind"] == "inline":
                set_paragraph_text(doc.paragraphs[f["para"]], f["name"] + "：" + str(val))
            else:
                continue
            written.append(f["name"])
        except Exception:
            continue
    return written


# ---------------------------------------------------------------- 列表区块
#
# 表单里除了「一个字段一个值」，还有成块的多行数据：学习经历、工作经历、
# 论文、家庭成员。模板里表现为「表头行 + 若干空行/示例行」，
# 资料里表现为「分节标题 + 表头行 + 数据行」。


def _row_all_empty_or_example(row):
    texts = _row_texts(row)
    if not any(texts):
        return True
    return all((not t) or is_example_text(t) for t in texts)


def analyze_list_tables(doc):
    """识别模板中的列表表。

    判定：首行是表头（≥2 个非空且较短的单元格），
    且其余所有行均为「整行空」或「全是示例文本」。
    """
    out = []
    for ti, tb in enumerate(doc.tables):
        if len(tb.rows) < 2:
            continue
        head = [t for t in _row_texts(tb.rows[0]) if t]
        if len(head) < 2 or any(len(t) > 16 for t in head):
            continue
        if not all(_row_all_empty_or_example(r) for r in tb.rows[1:]):
            continue
        out.append({
            "table": ti,
            "header": [norm_label(t) for t in _row_texts(tb.rows[0])],
            "data_rows": list(range(1, len(tb.rows))),
        })
    return out


def extract_list_sections(doc):
    """从资料文档抽取列表区块。

    返回 {分节标题: {"header": [...], "rows": [[...], ...]}}
    """
    sections = {}
    for tb in doc.tables:
        rows = tb.rows
        i = 0
        while i < len(rows):
            if not _is_section_title(rows[i]):
                i += 1
                continue
            title_texts = [t for t in _row_texts(rows[i]) if t]
            title = norm_label(title_texts[0]) if title_texts else ""
            if not title or i + 1 >= len(rows):
                i += 1
                continue
            header = [norm_label(t) for t in _row_texts(rows[i + 1])]
            if len([h for h in header if h]) < 2:
                i += 1
                continue
            data, j = [], i + 2
            while j < len(rows):
                if _is_section_title(rows[j]):
                    break
                texts = _row_texts(rows[j])
                # 列数必须与表头一致，否则不是数据行。
                # 学习经历末尾的「最高学位论文标题及内容简介」只有 2 列，
                # 不筛掉就会被当成一行学习经历写进模板。
                if any(texts) and len(texts) == len(header):
                    data.append(texts)
                j += 1
            if data:
                sections[title] = {"header": header, "rows": data}
            i = j
    return sections


# 需要把资料里多列拼起来的模板列。
# 例如模板要「学校、学院及专业」一栏，资料里院校和专业是分开的两列。
_CONCAT_COLUMNS = {
    "学校、学院及专业": ["毕业院校", "所学专业"],
    "毕业院校及专业": ["毕业院校", "所学专业"],
    "工作单位及职务": ["工作单位", "职务"],
}


def match_columns_multi(tpl_header, src_header, threshold=0.3):
    """match_columns 的增强版：允许一个模板列对应资料的多个列。

    返回列表，元素为 `int`（单列）、`list[int]`（多列合并，按顺序拼接）或 None。
    """
    single = match_columns(tpl_header, src_header, threshold)
    out = []
    for ti, tname in enumerate(tpl_header):
        parts = _CONCAT_COLUMNS.get(tname)
        if parts:
            idxs = []
            for p in parts:                    # 按别名表顺序取，保证"院校在前、专业在后"
                for si, s in enumerate(src_header):
                    if s and (p == s or p in s or s in p) and si not in idxs:
                        idxs.append(si)
                        break
            if idxs:
                out.append(idxs)
                continue
        out.append(single[ti])
    return out


def _col_sim(a, b):
    """列名相似度 0..1：先整串包含，再按字符重叠"""
    if not a or not b:
        return 0.0
    if a == b:
        return 1.0
    if a in b or b in a:
        return 0.9
    return len(set(a) & set(b)) / max(len(set(a)), len(set(b)))


# 常见人事/申报表格的列名别名。
# 有些对应关系字面完全不像——「工作职位」对「职务」、「所获学位」对「学习阶段」——
# 纯字符相似度算不出来（相似度只有 0.25），只能靠别名表。
_COLUMN_ALIASES = {
    "起止时间": ["起止时间", "起止", "时间"],
    "所在单位": ["工作单位", "单位", "所在单位"],
    "工作职位": ["职务", "职位", "岗位"],
    "所获学位": ["学习阶段", "学位"],
    "学历学位": ["学习阶段", "学位", "学历"],
    "学校、学院及专业": ["毕业院校", "所学专业", "学校", "院校", "专业"],
    "毕业院校及专业": ["毕业院校", "所学专业", "学校", "院校", "专业"],
    "作者排序": ["本人排名", "作者排序", "排名"],
    "期刊名称": ["刊物名称", "期刊", "刊物"],
    "出版号（ISSN）": ["ISSN号", "ISSN", "出版号"],
    "工作单位及职务": ["工作单位", "职务"],
    "与本人关系": ["与本人关系", "关系"],
    "出生年月": ["出生年月", "出生日期"],
    "论文名称": ["论文名称", "论文题目", "题目"],
    "联系电话": ["联系电话", "电话", "联系方式"],
    "个人邮箱": ["电子邮箱", "邮箱", "电子邮件"],
    "身份证号": ["身份证号码", "身份证号", "身份证"],
}


def match_columns(tpl_header, src_header, threshold=0.3):
    """模板列 -> 资料列下标，一对一的贪心匹配。

    两步：先查别名表（精确/包含），再用字符相似度兜底。
    **每列资料最多被用一次**——不做这个约束的话，
    「所在单位」和「工作职位」会同时匹配到资料里的「工作单位」。
    """
    n_t = len(tpl_header)

    def score(t, s):
        if not t or not s:
            return 0.0
        for idx, a in enumerate(_COLUMN_ALIASES.get(t, [])):
            if a == s:
                return 1.0
            if a and (a in s or s in a):
                # 越靠前的别名越优先（别名表按「更贴切」排序），
                # 减去一点分数保证排序时前面的赢
                return 0.95 - idx * 0.01
        return _col_sim(t, s)

    pairs = []
    for ti, t in enumerate(tpl_header):
        for si, s in enumerate(src_header):
            sc = score(t, s)
            if sc >= threshold:
                pairs.append((sc, ti, si))
    pairs.sort(reverse=True)

    mapping = [None] * n_t
    used = set()
    for sc, ti, si in pairs:
        if mapping[ti] is None and si not in used:
            mapping[ti] = si
            used.add(si)
    return mapping


def fill_list_table(doc, table_index, header, data_rows,
                    src_rows, col_map, append_rows=True):
    """把资料的多行数据填进模板的列表表。

    col_map: 模板列 -> 资料列下标（可用 match_columns 生成，也可手工指定）
    append_rows: 数据行多于模板空行时，复制最后一行追加
    返回实际写入的行数。
    """
    def _cell_value(src_row, cm):
        """cm 可以是单个下标，也可以是下标列表（多列合并）"""
        if cm is None:
            return ""
        if isinstance(cm, (list, tuple)):
            parts = [src_row[i].strip() for i in cm if i < len(src_row)]
            return " ".join(p for p in parts if p)
        return src_row[cm].strip() if cm < len(src_row) else ""

    tb = doc.tables[table_index]
    written = 0

    # 过滤空数据行：映射列上非空值少于 2 个的行丢掉。
    # 学习经历里的「中专」「大专」只有阶段名、其余全空，
    # 不过滤就会占掉模板里的一行，还会把示例文本留在旁边。
    usable = [r for r in src_rows
              if sum(1 for cm in col_map if _cell_value(r, cm)) >= 2]
    src_rows = usable

    for k, src_row in enumerate(src_rows):
        if k >= len(data_rows):
            if not append_rows:
                break
            # 复制最后一行再追加
            last = tb.rows[data_rows[-1]]._tr
            new_tr = copy.deepcopy(last)
            last.addnext(new_tr)
            data_rows.append(len(tb.rows) - 1)
        ri = data_rows[k]
        row = tb.rows[ri]
        cells = iter_row_cells(row)
        for ci, (_gci, cell) in enumerate(cells):
            cm = col_map[ci] if ci < len(col_map) else None
            val = _cell_value(src_row, cm)
            if val:
                set_cell_text(cell, val)
                written += 1
            elif cm is not None and is_example_text(cell.text):
                # 没有值可填，但格子里是示例文本 → 清掉，否则会留下 "XXX大学XXX学院"
                set_cell_text(cell, "")

    # 模板里剩下的多余空行（数据行少于模板行数）也要清掉示例文本
    for ri in data_rows[len(src_rows):]:
        for _gci, cell in iter_row_cells(tb.rows[ri]):
            if is_example_text(cell.text):
                set_cell_text(cell, "")
    return written


# ---------------------------------------------------------------- 论文引用解析
#
# 简历里「发表论文」往往是一段 APA 引用文本（常放在文本框里），而不是表格。
# extract_list_sections 只读表格，读不到这些文本段落，导致论文的 ISSN、
# 影响因子、收录、年卷期页等信息被白白丢弃。这里把文本引用解析成结构化字段，
# 再转成与列表区块同构的 {分节标题: {header, rows}}，供 match_columns +
# fill_list_table 直接复用。

# 论文表的标准列（与常见申报书「发表论文情况」表头对齐，不含「图书馆认定意见」空列）
PAPER_HEADER = ["作者排序", "论文名称", "期刊名称", "出版号（ISSN）",
                "影响因子", "收录情况", "是否通讯作者", "年、卷（期）：页"]

# 常见经济/管理类期刊 -> ISSN（打印版）。联网查询需要稳定权威来源且可能
# 被反爬/限流，先内置高频期刊；表外期刊留空由用户补。
_JOURNAL_ISSN = {
    "journal of economic behavior and organization": "0167-2681",
    "post communist economies": "1463-1377",
    "post-communist economies": "1463-1377",
    "american economic review": "0002-8282",
    "quarterly journal of economics": "0033-5533",
    "journal of political economy": "0022-3808",
    "econometrica": "0012-9682",
    "review of economic studies": "0034-6527",
    "review of economics and statistics": "0034-6535",
    "journal of econometrics": "0304-4076",
    "economic journal": "0013-0133",
    "journal of development economics": "0304-3878",
    "journal of public economics": "0047-2727",
    "journal of finance": "0022-1082",
    "review of financial studies": "0893-9454",
    "journal of monetary economics": "0304-3932",
    "european economic review": "0014-2921",
    "journal of comparative economics": "0147-5967",
    "china economic review": "1043-951X",
}

_CN_DIGITS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
              "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _journal_key(journal):
    """期刊名归一化：小写、&→and、非字母→空格、压缩空格，便于查 ISSN 表。"""
    j = (journal or "").lower().replace("&", " and ")
    j = re.sub(r"[^a-z]+", " ", j)
    return " ".join(j.split())


def _rank_name(idx):
    """作者下标 -> 中文排名：0->第一作者，1->第二作者……"""
    cn = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if 0 <= idx < len(cn):
        return "第%s作者" % cn[idx]
    return "第%d作者" % (idx + 1)


# 英文作者：'Lee, H.' / 'Zou, F.'；中文作者由 _split_authors 兜底
_AUTHOR_RE = re.compile(
    r"([A-Z][A-Za-z'\-]+)\s*,?\s*([A-Z](?:\s*\.\s*[A-Z])*\.?)")


def _split_authors(s):
    """'Lee, H., Zhao, K. and Zou, F.' -> ['Lee, H.', 'Zhao, K.', 'Zou, F.']"""
    s = re.sub(r"\s+(?:and|&)\s+", ", ", s)
    s = s.replace("，", ", ").replace("、", ", ")
    pairs = _AUTHOR_RE.findall(s)
    if pairs:
        return ["%s, %s" % (last, init.strip()) for last, init in pairs]
    return [a.strip() for a in re.split(r"[,、]", s) if a.strip()]


def _find_self_index(authors, hints):
    """在作者列表里定位「本人」：hints 里任一关键词命中某作者即返回其下标。"""
    for hint in hints or []:
        h = (hint or "").strip().lower()
        if not h:
            continue
        for i, a in enumerate(authors):
            if h in a.lower():
                return i
    return None


def _fmt_vol_pages(year, vol, issue, pages):
    if vol and pages:
        if issue:
            return "%s, %s(%s): %s" % (year, vol, issue, pages)
        return "%s, %s: %s" % (year, vol, pages)
    if pages:
        return "%s: %s" % (year, pages)
    return year or ""


def parse_citation(text, self_hints=()):
    """解析一条 APA 风格论文引用，返回结构化 dict；解析不出返回 None。

    兼容形如：
        Lee, H., Zhao, K. and Zou, F. (2022) "Does ...?" Journal of Economic
        Behavior and Organization. (196) 330-345. SSCI Q2，中科院3区，
        影响因子2.6. 文章引用10次
    """
    t = (text or "").strip()
    if not t:
        return None
    m = re.match(r"^(.*?)\((\d{4})\)", t)
    if not m:
        return None
    author_str, year = m.group(1), m.group(2)
    rest = t[m.end():]
    authors = _split_authors(author_str)
    info = {"authors": authors, "year": year, "raw": t}

    m_title = re.search(r"[“\"](.+?)[”\"]", rest)
    if m_title:
        info["title"] = m_title.group(1).strip()
        after_title = rest[m_title.end():]
    else:
        after_title = rest

    # 年卷期页：先定位「页码范围」，再往前找可选的「卷(期)」
    vol = issue = pages = None
    m_pages = re.search(r"(\d{1,4})\s*[-–—]\s*(\d{1,4})", after_title)
    if m_pages:
        pages = "%s-%s" % (m_pages.group(1), m_pages.group(2))
        pre = after_title[:m_pages.start()]
        m_vol = re.search(
            r"\(?\s*(\d{1,4})\s*\)?\s*(?:\(\s*(\d{1,3})\s*\))?\s*[:：]?\s*$", pre)
        if m_vol and m_vol.group(1):
            vol, issue = m_vol.group(1), m_vol.group(2)
            journal_part = pre[:m_vol.start()]
        else:
            journal_part = pre
    else:
        journal_part = after_title

    journal = re.sub(r"[.,;:，。；：、\s]+$", "", journal_part).strip()
    info["journal"] = journal
    info["volume"], info["issue"], info["pages"] = vol, issue, pages
    info["year_vol_pages"] = _fmt_vol_pages(year, vol, issue, pages)

    # 收录情况：SSCI/SCI + JCR 分区 Qx + 中科院分区
    indexing = []
    up = rest.upper()
    if "SSCI" in up:
        indexing.append("SSCI")
    elif re.search(r"\bSCI\b", up):
        indexing.append("SCI")
    m_q = re.search(r"Q([1-4])", up)
    if m_q:
        indexing.append("Q%s" % m_q.group(1))
    m_cas = re.search(r"中科院\s*([1-4])\s*区", rest)
    if m_cas:
        indexing.append("中科院%s区" % m_cas.group(1))
    info["indexing"] = "，".join(indexing)

    # 影响因子
    m_if = (re.search(r"影响因子\s*(\d+(?:\.\d+)?)", rest)
            or re.search(r"\bIF\s*[:=]?\s*(\d+(?:\.\d+)?)", rest, re.I))
    if m_if:
        info["impact_factor"] = m_if.group(1)

    # 引用次数
    m_cit = re.search(r"(?:引用|被引|cited)\s*([\d]+)\s*次", rest, re.I)
    if m_cit:
        info["cited"] = m_cit.group(1)

    # 作者排序：据本人在作者列表中的位置推导
    idx = _find_self_index(authors, self_hints)
    info["self_index"] = idx
    info["rank"] = _rank_name(idx) if idx is not None else ""

    # 是否通讯作者：引用文本里出现「通讯作者」字样或星号/† 标记才判「是」
    info["corresponding"] = "是" if re.search(r"通讯作者|通迅作者", rest) else ""

    # ISSN：按期刊名查内置表
    info["issn"] = _JOURNAL_ISSN.get(_journal_key(journal), "")

    return info


def extract_papers_from_text(text, self_hints=()):
    """从资料文档的全文文本里解析「发表论文」章节。

    返回 {分节标题: {"header": PAPER_HEADER, "rows": [[...]]}}，
    供 match_columns + fill_list_table 复用；没有论文时返回 {}。
    """
    if not text:
        return {}
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    # 定位「发表论文」章节：从该标题到下一个同级标题（工作论文/课题/…）之间
    start = None
    for i, ln in enumerate(lines):
        if re.match(r"^发表论文|^已发表论文|^发表的论文", ln):
            start = i
            break
    if start is None:
        return {}
    end = len(lines)
    for j in range(start + 1, len(lines)):
        if re.match(r"^(工作论文|工作中论文|在投论文|课题|研究成果|教学|学术交流|"
                    r"获奖|专利|专著|教材|主持|参与)", lines[j]):
            end = j
            break

    rows = []
    for ln in lines[start + 1:end]:
        info = parse_citation(ln, self_hints)
        if not info:
            continue
        rows.append([
            info.get("rank", ""),
            info.get("title", ""),
            info.get("journal", ""),
            info.get("issn", ""),
            info.get("impact_factor", ""),
            info.get("indexing", ""),
            info.get("corresponding", ""),
            info.get("year_vol_pages", ""),
        ])
    if not rows:
        return {}
    return {"发表论文": {"header": PAPER_HEADER, "rows": rows}}
