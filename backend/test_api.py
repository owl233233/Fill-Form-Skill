"""API 端到端测试：完整用户流程"""
import json
import sys

import requests

BASE = "http://localhost:8765"
# 原为云端硬编码路径 /workspace/form-filler/sample 与 /tmp/ff_test，本地不可用，改为相对定位
import os
_PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SAMPLE = os.path.join(_PROJECT_DIR, "sample")
TMP = os.path.join(_PROJECT_DIR, "tmp", "ff_test")
os.makedirs(TMP, exist_ok=True)
ok, fail = 0, 0

def check(name, cond, detail=""):
    global ok, fail
    if cond:
        ok += 1
        print(f"  ✓ {name}")
    else:
        fail += 1
        print(f"  ✗ {name}  {detail}")

print("== 1. 上传资料文件 ==")
for fn in ["李云飞个人简历.docx", "岗位信息.txt", "头像.png"]:
    with open(f"{SAMPLE}/{fn}", "rb") as f:
        r = requests.post(f"{BASE}/api/library/upload", files={"file": (fn, f)})
    check(f"上传 {fn}", r.status_code == 200, r.text[:200])
    if r.status_code == 200:
        d = r.json()
        if fn.endswith((".docx", ".txt")):
            print(f"    提取 {d['extracted_count']} 条: {d.get('added_keys', [])}")

lib = requests.get(f"{BASE}/api/library").json()
print(f"  资料库: {len(lib['entries'])} 条信息 / {len(lib['files'])} 个文件 / {len(lib['images'])} 张图片")
check("信息条目 ≥ 12", len(lib["entries"]) >= 12, str(len(lib["entries"])))
check("图片入库", len(lib["images"]) == 1)

print("== 2. 上传模板（带命名规则） ==")
with open(f"{SAMPLE}/员工入职登记表模板.docx", "rb") as f:
    r = requests.post(f"{BASE}/api/templates/upload",
                      files={"file": ("员工入职登记表模板.docx", f)},
                      data={"naming_pattern": "{{姓名}}_入职登记表_{{当前日期}}"})
check("上传 docx 模板", r.status_code == 200, r.text[:300])
tpl_docx = r.json()
print(f"    占位符: {tpl_docx['text_fields']}  图片: {tpl_docx['img_fields']}")
check("识别 16 个文本占位符", len(tpl_docx["text_fields"]) == 16)
check("识别图片占位符", tpl_docx["img_fields"] == ["头像"])

with open(f"{SAMPLE}/产品报价单模板.xlsx", "rb") as f:
    r = requests.post(f"{BASE}/api/templates/upload",
                      files={"file": ("产品报价单模板.xlsx", f)},
                      data={"naming_pattern": "{{公司}}_报价单_{{姓名}}"})
check("上传 xlsx 模板", r.status_code == 200, r.text[:300])
tpl_xlsx = r.json()

print("== 3. 预览填充（缺失检测） ==")
r = requests.post(f"{BASE}/api/fill/preview", data={"template_id": tpl_docx["id"]})
check("preview 成功", r.status_code == 200, r.text[:300])
pv = r.json()
missing = [f["name"] for f in pv["fields"] if f["missing"]]
matched = {f["name"]: f["value"] for f in pv["fields"] if f["value"]}
print(f"    已匹配 {len(matched)} 项; 缺失: {missing}")
print(f"    图片匹配: {pv['images']}")
print(f"    建议文件名: {pv['suggested_filename']}")
check("缺失检测正确(签名/签署日期)", set(missing) == {"签名", "签署日期"}, str(missing))
check("智能匹配 地址←现居地址", "地址" in matched and "上海" in matched["地址"])
check("智能匹配 部门←入职部门", "部门" in matched and matched["部门"] == "人工智能事业部")
check("编号←学号 别名匹配", matched.get("编号") == "ZC-2026-0889")
check("动态字段当前日期不缺失", all(not f["missing"] for f in pv["fields"] if f["name"] == "当前日期"))
check("图片头像自动匹配", pv["images"][0]["matched"] is not None)
check("命名规则预渲染", "李云飞" in pv["suggested_filename"], pv["suggested_filename"])

print("== 4. 缺失时生成 → 应返回 422 询问 ==")
r = requests.post(f"{BASE}/api/fill/generate",
                  data={"template_id": tpl_docx["id"], "values": json.dumps({})})
check("缺失时返回 422", r.status_code == 422, f"{r.status_code} {r.text[:200]}")
check("返回缺失字段列表", set(r.json().get("missing", [])) == {"签名", "签署日期"})

print("== 5. 补充缺失信息后生成 ==")
values = {"签名": "李云飞", "签署日期": "2026-08-29"}
r = requests.post(f"{BASE}/api/fill/generate",
                  data={"template_id": tpl_docx["id"],
                        "values": json.dumps(values),
                        "filename": pv["suggested_filename"],
                        "save_to_library": "true"})
check("生成成功", r.status_code == 200, r.text[:300])
gen = r.json()
print(f"    文件名: {gen['filename']}  图片: {gen['images_inserted']}  存库: {gen['saved_keys']}")
check("按命名规则命名", gen["filename"].startswith("李云飞_入职登记表_2026"), gen["filename"])
check("图片已插入", gen["images_inserted"] == 1)
check("补充信息已存库", "签名" in gen["saved_keys"])

# 下载验证
r = requests.get(BASE + gen["download_url"])
check("文件可下载", r.status_code == 200 and len(r.content) > 5000)
_out_docx = os.path.join(TMP, "api_output.docx")
with open(_out_docx, "wb") as f:
    f.write(r.content)
from docx import Document
d = Document(_out_docx)
all_text = "\n".join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += "\n" + cell.text
check("API生成文件无残留占位符", "{{" not in all_text)
check("API生成文件含签名", "李云飞" in all_text)
check("API生成文件含图片", len(d.inline_shapes) == 1)

lib2 = requests.get(f"{BASE}/api/library").json()
check("资料库新增签名条目", any(e["key"] == "签名" for e in lib2["entries"]))

print("== 6. Excel 模板流程 ==")
r = requests.post(f"{BASE}/api/fill/preview", data={"template_id": tpl_xlsx["id"]})
pvx = r.json()
missing_x = [f["name"] for f in pvx["fields"] if f["missing"]]
print(f"    xlsx 缺失: {missing_x}; 建议名: {pvx['suggested_filename']}")
vals_x = {f: "示例值" for f in missing_x}
vals_x.update({"产品名称": "智能填表系统企业版", "数量": "10", "单价": "8800", "备注": "含一年技术支持"})
r = requests.post(f"{BASE}/api/fill/generate",
                  data={"template_id": tpl_xlsx["id"],
                        "values": json.dumps(vals_x),
                        "filename": pvx["suggested_filename"]})
check("xlsx 生成成功", r.status_code == 200, r.text[:300])
genx = r.json()
print(f"    文件名: {genx['filename']}")
check("xlsx 按命名规则命名", genx["filename"].startswith("智创未来科技有限公司_报价单"), genx["filename"])
r = requests.get(BASE + genx["download_url"])
_out_xlsx = os.path.join(TMP, "api_output.xlsx")
with open(_out_xlsx, "wb") as f:
    f.write(r.content)
from openpyxl import load_workbook
wb = load_workbook(_out_xlsx)
ws = wb.active
cells = " ".join(str(c.value) for row in ws.iter_rows() for c in row if c.value)
check("xlsx 无残留占位符", "{{" not in cells)

print("== 7. 生成记录 ==")
recs = requests.get(f"{BASE}/api/records").json()
check("记录 2 条", len(recs) == 2, str(len(recs)))
check("记录含文件名", any(r["filename"] == gen["filename"] for r in recs))

print("== 8. 前端页面 ==")
r = requests.get(BASE + "/")
check("首页 200", r.status_code == 200 and "智填" in r.text)
r = requests.get(BASE + "/app.js")
check("静态资源正常", r.status_code == 200)

print(f"\n结果：通过 {ok} 项，失败 {fail} 项")
sys.exit(1 if fail else 0)
