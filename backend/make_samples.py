"""构造测试样例：带格式的 docx/xlsx 模板 + 资料文件 + 图片素材"""
import os
import sys

# 原为云端硬编码路径 /workspace/form-filler/backend，本地不可用，改为相对定位
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from PIL import Image, ImageDraw

SAMPLE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "sample")
os.makedirs(SAMPLE_DIR, exist_ok=True)

# ---------------- 1. 带格式的 Word 模板 ----------------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "DengXian"
style.font.size = Pt(11)

# 标题：居中、大号、深蓝色
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("员工入职登记表")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8F)
run.font.name = "DengXian"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("填表日期：{{当前日期}}　　编号：{{编号}}")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# 说明段：斜体
note = doc.add_paragraph()
r = note.add_run("说明：本表由人力资源部存档，请如实填写。")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 信息表格
table = doc.add_table(rows=6, cols=4)
table.style = "Table Grid"
rows_data = [
    ("姓名", "{{姓名}}", "性别", "{{性别}}"),
    ("出生日期", "{{出生日期}}", "联系电话", "{{电话}}"),
    ("电子邮箱", "{{邮箱}}", "紧急联系人", "{{紧急联系人}}"),
    ("现居地址", "{{地址}}", "", ""),
    ("毕业院校", "{{学校}}", "专业", "{{专业}}"),
    ("入职部门", "{{部门}}", "职位", "{{职位}}"),
]
for i, (k1, v1, k2, v2) in enumerate(rows_data):
    cells = table.rows[i].cells
    for j, txt in enumerate((k1, v1, k2, v2)):
        cells[j].text = txt
        for p in cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10.5)
                if j % 2 == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# 图片占位符 + 签名段
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("员工照片：{{img:头像}}")
r.font.size = Pt(10)

sig = doc.add_paragraph()
r = sig.add_run("本人承诺以上信息真实有效。")
r.font.size = Pt(10.5)
sig2 = doc.add_paragraph()
r = sig2.add_run("员工签名：{{签名}}　　　　日期：{{签署日期}}")
r.font.size = Pt(10.5)

# 页眉
sec = doc.sections[0]
hp = sec.header.paragraphs[0]
hp.text = "{{公司}} · 内部资料"
for r in hp.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

tpl_docx = os.path.join(SAMPLE_DIR, "员工入职登记表模板.docx")
doc.save(tpl_docx)
print("✓ Word 模板:", tpl_docx)

# ---------------- 2. 带格式的 Excel 模板 ----------------
wb = Workbook()
ws = wb.active
ws.title = "报价单"

thin = Side(style="thin", color="BFBFBF")
border = Border(left=thin, right=thin, top=thin, bottom=thin)
head_fill = PatternFill("solid", fgColor="1F3A8F")

ws.merge_cells("A1:D1")
c = ws["A1"]
c.value = "{{公司}}产品报价单"
c.font = Font(size=16, bold=True, color="FFFFFF")
c.fill = head_fill
c.alignment = Alignment(horizontal="center", vertical="center")
ws.row_dimensions[1].height = 30

meta = [("客户名称：", "{{姓名}}", "联系电话：", "{{电话}}"),
        ("报价日期：", "{{当前日期}}", "电子邮箱：", "{{邮箱}}")]
for i, row in enumerate(meta, start=2):
    for j, v in enumerate(row):
        cell = ws.cell(row=i, column=j + 1, value=v)
        cell.font = Font(size=10.5, bold=(j % 2 == 0))
        if j % 2 == 0:
            cell.fill = PatternFill("solid", fgColor="EEF2FF")
        cell.border = border

headers = ["产品名称", "数量", "单价（元）", "备注"]
for j, hname in enumerate(headers, start=1):
    cell = ws.cell(row=5, column=j, value=hname)
    cell.font = Font(bold=True, color="FFFFFF")
    cell.fill = PatternFill("solid", fgColor="4F46E5")
    cell.alignment = Alignment(horizontal="center")
    cell.border = border

items = [("{{产品名称}}", "{{数量}}", "{{单价}}", "{{备注}}")]
for i, row in enumerate(items, start=6):
    for j, v in enumerate(row, start=1):
        cell = ws.cell(row=i, column=j, value=v)
        cell.border = border
        cell.font = Font(size=10.5)

for col, w in zip("ABCD", (22, 12, 14, 20)):
    ws.column_dimensions[col].width = w

tpl_xlsx = os.path.join(SAMPLE_DIR, "产品报价单模板.xlsx")
wb.save(tpl_xlsx)
print("✓ Excel 模板:", tpl_xlsx)

# ---------------- 3. 资料文件（模拟用户上传的个人简历 docx） ----------------
d2 = Document()
d2.add_heading("个人简历", level=1)
lines = [
    "姓名：李云飞",
    "性别：男",
    "出生日期：1995-06-18",
    "联系电话：13812345678",
    "电子邮箱：liyunfei@example.com",
    "现居地址：上海市浦东新区张江高科技园区博云路2号",
    "毕业院校：复旦大学",
    "专业：计算机科学与技术",
    "紧急联系人：李建国（父亲）13900001111",
]
for line in lines:
    d2.add_paragraph(line)
resume = os.path.join(SAMPLE_DIR, "李云飞个人简历.docx")
d2.save(resume)
print("✓ 资料文件:", resume)

# ---------------- 4. 补充资料 txt（含部门职位等） ----------------
extra = os.path.join(SAMPLE_DIR, "岗位信息.txt")
with open(extra, "w", encoding="utf-8") as f:
    f.write("入职部门：人工智能事业部\n职位：算法工程师\n公司：智创未来科技有限公司\n编号：ZC-2026-0889\n")
print("✓ 资料文件:", extra)

# ---------------- 5. 图片素材（头像） ----------------
img = Image.new("RGB", (240, 320), "#4F46E5")
draw = ImageDraw.Draw(img)
draw.ellipse([70, 60, 170, 160], fill="#EEF2FF")
draw.rectangle([40, 190, 200, 300], fill="#EEF2FF")
avatar = os.path.join(SAMPLE_DIR, "头像.png")
img.save(avatar)
print("✓ 图片素材:", avatar)

print("\n全部样例生成完毕")
